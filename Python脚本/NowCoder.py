import time
import docx
from bs4 import BeautifulSoup
import requests
from selenium import webdriver
urlbase = "https://www.nowcoder.com"
doc = docx.Document()
def getIndexPage(url):
    #chorm_driver = 'D:\Download\Google\chromedriver_win32\chromedriver.exe'
    driver = webdriver.Chrome()
    #这个是牛客网java实习面经的界面，可以根据自己的需要，找好网页，然后将连接复制到这里来
    driver.get(url)
    time.sleep(3)
    js = "return action=document.body.scrollHeight"
    height = driver.execute_script(js)
    driver.execute_script('window.scrollTo(0, document.body.scrollHeight)')
    time.sleep(5)
    t1 = int(time.time())
    status = True
    num = 0
    #这里的一堆代码就是将滚动条拉到最下面，让资源加载完毕。
    while status:
        t2 = int(time.time())
        if t2 - t1 < 30:
            new_height = driver.execute_script(js)
            if new_height > height:
                time.sleep(1)
                driver.execute_script('window.scrollTo(0, document.body.scrollHeight)')
                height = new_height
                t1 = int(time.time())
        elif num < 3:
            time.sleep(3)
            num = num + 1
        else:
            print("滚动条已经处于页面最下方！")
            status = False
            driver.execute_script('window.scrollTo(0, 0)')
            break
    content = driver.page_source
    return content

def getUrl(page):
    soup = BeautifulSoup(page, 'lxml')
    list=[]
    for ul in soup.select(".js-nc-wrap-link"):
        list.append(ul.attrs['data-href'])
    return list

def getPageDetail(urll):
    try:
        response = requests.get(urll)
        if response.status_code == 200:
            return response.text
        return None
    except ConnectionError:
        print('Error occurred')
        return None
def parseDetail(page):
    soup = BeautifulSoup(page, 'lxml')
    return soup.select(".discuss-title")[0].get_text()+soup.select(".post-time")[0].get_text()+soup.select(".post-topic-des")[0].get_text()
def pageSave(str,count):
    doc.add_paragraph("**********************************第%d篇*************************************" % count)
    doc.add_paragraph(str)
def main():
    page = getIndexPage("https://www.nowcoder.com/discuss/experience?tagId=639&order=3&companyId=652&phaseId=3")
    list = getUrl(page)
    print("一共有%d篇"%len(list))
    count = 0
    for item in list:
        content = getPageDetail(urlbase+item)
        str = parseDetail(content)
        pageSave(str,count)
        count = count + 1
        print("进行到第%d篇了,Url为:" % count+urlbase+item)
        if count>500:
            break
    # str = getPageDetail("https://www.nowcoder.com/discuss/398783")
    # s = parseDetail(str)
    # doc.add_paragraph(s)
    doc.save("截止2020-08-28日牛客Java工程师滴滴社招面经合集.docx")


if __name__ == '__main__':
    main()
